import os
from openai import OpenAI
import pymupdf
import pandas as pd
import numpy as np
import json
import io
import re
from PIL import Image
import pytesseract
from dotenv import load_dotenv
from pydantic import BaseModel
from typing import List, Optional
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

from pylatex import NoEscape, Command
from pylatex import Document as DocPy
from pylatex.utils import escape_latex

#brew install tesseract

load_dotenv()

client = OpenAI(
  api_key=os.environ.get("OPENAI_KEY")
)

#Text Processing

def process_file(path: str) -> str:
    """
    processes file to text. supports PDF and image files (PNG, JPG, JPEG).
    """
    if(path.lower().endswith(".pdf")):
        return get_text_from_pdf(path)
    elif(path.lower().endswith(".docx")):
        return get_text_from_docx(path)
    elif(path.lower().endswith((".png", ".jpg", ".jpeg"))):
        return get_text_from_img(path)
    else:
        raise ValueError("Unsupported file type. Only PDFs, DOCXs, and images are supported.")

def get_text_from_pdf(path: str) -> str:
    """
    Uses pymupdf to extract text from PDF
    """
    doc = pymupdf.open(path)
    text = "\n".join(page.get_text() for page in doc)
    return text.strip()

def get_text_from_img(path: str) -> str:
    """
    Uses pytesseract to extract text from PDF
    """
    image = Image.open(path)
    text = pytesseract.image_to_string(image)
    return text.strip()

def get_text_from_docx(path: str) -> str:
    """
    Uses Python-docx to extract text from .DOCX
    """
    document = Document(path)
    text_chunks = []

    for paragraph in document.paragraphs:
        if paragraph.text.split():
            text_chunks.append(paragraph.text)

    for table in document.tables:
        for row in table.rows:
            row_text = "\t".join(cell.text.strip() for cell in row.cells)
            text_chunks.append(row_text)

    return "\n".join(text_chunks).strip()

#Prompt Engineering

#LLM Output Structure
class Metadata(BaseModel):
    title: Optional[str] = None
    date: Optional[str] = None
    author: Optional[str] = None
    course: Optional[str] = None
    additional_info: Optional[str] = None

class Subquestion(BaseModel):
    question: str
    answer: str

class QuestionAnswer(BaseModel):
    question: str
    answer: str
    subquestions: Optional[List[Subquestion]] = None

class OrganizedOutput(BaseModel):
    metadata: Metadata
    content: List[QuestionAnswer]

def organize_LLM(questions_text: str, answers_text: str, model:str ="gpt-4o") -> dict:
    """
    Uses OpenAI's (GPT-4o) LLM Model to grammar-check & intelligently organize the questions and answers together into a JSON.
    """

    prompt = (
        "You are an intelligent assistant. Your task is to correct spelling errors, organize, "
        "and accurately pair questions with answers in a structured JSON format. Follow these **strict rules**:\n"
        
        "1. If metadata is provided (e.g., title, date, author, course, etc.), include it under a 'metadata' key at the top.\n"
        "2. Place all questions and answers under a 'content' key.\n"
        "3. Use the provided numbering system (e.g., 1, 1a, 1b). Assign sequential numbers starting from 1 if none are given. Include an answer explanation (semi-concise) if given.\n"
        "4. Handling main questions with subparts:\n"
        "   a. **If the main question is answered as a whole**, include the entire question (with subparts) under a single 'question' key and give a single answer.\n"
        "   b. **If subparts are answered separately**, list the main question if it has an answer, and include subparts under 'subquestions', each with 'question' and 'answer' keys.\n"
        "5. Use this concise format for each entry:\n"
        "   { \"<number>\": { \"question\": \"<question_text>\", \"answer\": \"<answer_text>\", \"subquestions\": [{ \"question\": \"<subquestion_text>\", \"answer\": \"<subanswer_text>\" }] } }\n"
        "   - Omit 'subquestions' if there are none.\n"
        "6. Use LaTeX for math content (e.g., $\\(x^2 + 4x + 4 = 0\\)$. Ensure math content is enclosed in $$).\n"
        "7. Keep spelling/grammar corrections minimal; assume the text is OCR-generated.\n"
        "8. Ensure the output is valid JSON without enclosing triple backticks, code blocks, or explanations.\n\n"
        
        "Questions:\n"
        f"{questions_text}\n\n"
        "Answers:\n"
        f"{answers_text}\n\n"
        "Return the output as valid JSON."
    )

    response = client.beta.chat.completions.parse(
        model=model,
        messages=[
            {"role": "system", "content": "You are a helpful assistant for processing questions and answers."},
            {"role": "user", "content": prompt}
        ],
        max_tokens=3000,
        temperature=0.0,
        response_format=OrganizedOutput
    )

    
    try:
        #Dict of QA objects
        response_dict = response.choices[0].message.parsed.model_dump()

        return response_dict
    except (json.JSONDecodeError, KeyError):
        print(response)
        raise ValueError("Failed to parse response from OpenAI GPT.")
    except Exception as e:
        print(e)
        print(response_dict)

#Export Options

def export_output(data: dict, output_format: str = "json", output_path: str = "output"):
    if output_format == "json":
        save_to_json(dict, f"{output_path}.json")
    elif output_format == "csv":
        pd.DataFrame(data["content"]).to_csv(f"{output_path}.csv", index=False)
    elif output_format == "docx":
        save_to_docx(data, f"{output_path}.docx")
    # elif output_format == "LaTeX":
    #     save_to_LaTeX(data, f"{output_path}.tex")
    else:
        raise ValueError("File Type Error")


def save_to_json(data: dict, output_path: str = "output"):
    """
    Saves a dict to a .json file at path 'output_path'
    """
    with open(output_path, "w") as file:
        json.dump(data, file, indent=4)
    print(f"Output saved to {output_path}")

def save_to_docx(data: dict, output_path: str = "output"):
    """
    Saves a dict to a .docx file at path 'output_path'
    """
    document = Document()
    #Formatting Header Stuff
    metadata = data.get("metadata", {})

    if "title" in metadata:
        title = document.add_heading(metadata["title"], level=0)
        title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    if metadata:
        metadata_paragraph = document.add_paragraph()
        if "course" in metadata:
            metadata_paragraph.add_run("Course: ").bold = True
            metadata_paragraph.add_run(metadata["course"]).italic = True
            metadata_paragraph.add_run("\n")
        if "author" in metadata:
            metadata_paragraph.add_run("Author: ").bold = True
            metadata_paragraph.add_run(metadata["author"]).italic = True
            metadata_paragraph.add_run("\n")
        if "date" in metadata:
            metadata_paragraph.add_run("Date: ").bold = True
            metadata_paragraph.add_run(metadata["date"]).italic = True
            metadata_paragraph.add_run("\n")
        if "additional_info" in metadata:
            metadata_paragraph.add_run("Additional Info: ").bold = True
            metadata_paragraph.add_run(metadata["additional_info"]).italic = True
            metadata_paragraph.add_run("\n")

    document.add_paragraph() # Post Metadata Spacing
    
    #Qs & As Section
    for idx, qa in enumerate(data["content"], start = 1):
        question = qa["question"]
        answer = qa.get("answer") #remember this can be None if there exists subquestions
        subquestions = qa.get("subquestions")

        #Add MainQ
        document.add_paragraph(f"  Q{idx}: {question}", style='List Bullet')
        if answer:
            document.add_paragraph(f"  A{idx}: {answer}")

        #Add SubQs
        if subquestions:
            for sub_idx, subquestion in enumerate(subquestions, start=1):
                sub_label = chr(96 + sub_idx)
                document.add_paragraph(f"  Q{idx}.{sub_label}: {subquestion['question']}", style='List Bullet')
                document.add_paragraph(f"  A{idx}.{sub_label}: {subquestion['answer']}")

        document.add_paragraph()
    
    #Line Spacing Stuff
    for paragraph in document.paragraphs:
        paragraph_format = paragraph.paragraph_format
        paragraph_format.space_after = Pt(8)
        paragraph_format.line_spacing = 1.5
        paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    document.save(output_path)
    print(f"Output saved to {output_path}")

def main():
    #PDF to JSON Example
    question_filepath = "./pdfs/math110-midterm-questions-071312.pdf"
    answers_filepath = "./pdfs/math110-midterm-answers-071312.pdf"

    #Image to JSON Example
    # question_filepath = "./pngs/oop_exam_questions.png"
    # answers_filepath = "./pngs/oop_exam_answers.png"
    try:
        questions_text = process_file(question_filepath)
        answers_text = process_file(answers_filepath)
    except Exception as e:
        print(f"Error processing files: {e}")
        return
    
    #print(f"Questions: {questions_text} \n\n Answers: {answers_text}")
    paired_results = organize_LLM(questions_text, answers_text)
    save_to_json(paired_results, "output2.json")


if __name__ == "__main__":
    main()